"""
FundedNext Dispute Document Generator
======================================
Runs locally OR via GitHub Actions (no device required).

GitHub Actions: fill in the form under Actions → Generate Dispute DOCX → Run workflow.
                Upload per-dispute files (journal.pdf, invoice.pdf, etc.) to the
                pdfs/ folder via the GitHub web UI first, then trigger the workflow.

Local usage:    pip install -r requirements.txt
                python generate_docx.py          (reads config.json)

pdfs/ folder layout
---------------------
  checkout_flow.pdf            required — text, committed to repo, never changes
  refund_cancellation.pdf      required — text, committed to repo, never changes
  terms_of_service.pdf         required — text, committed to repo, never changes
  journal.pdf                  required — MT5 Manager export, upload per dispute
  invoice.pdf  or .png         optional — upload per dispute
  intercom.png or .pdf         optional — upload per dispute
  gmail.png    or .pdf         optional — upload per dispute
  admin.png    or .pdf         optional — upload per dispute
  drawdown_chart.png           optional — upload per dispute
"""

import os, re, json
from pathlib import Path
from io import BytesIO

import pdfplumber
import fitz
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
PDFS = BASE / "pdfs"

# ── Config ────────────────────────────────────────────────────────────────────
DEFAULTS = {
    "customer_name":  "Customer Name",
    "customer_email": "customer@example.com",
    "arn":            "000000000000000",
    "amount":         "0.00",
    "card_type":      "Visa",
    "card_number":    "**** **** **** ****",
    "purchase_date":  "2026-01-01",
    "account_login":  "N/A",
    "merchant_ref":   "N/A",
    "descriptor":     "FUNDEDNEXT*CHALLENGE",
    "ip":             "N/A",
    "ip_location":    "N/A",
    "access_time":    "N/A",
    "device":         "N/A",
    "device_id":      "N/A",
}

ENV_MAP = {
    "customer_name":  "CUSTOMER_NAME",
    "customer_email": "CUSTOMER_EMAIL",
    "arn":            "ARN",
    "amount":         "AMOUNT",
    "card_type":      "CARD_TYPE",
    "card_number":    "CARD_NUMBER",
    "purchase_date":  "PURCHASE_DATE",
    "account_login":  "ACCOUNT_LOGIN",
    "merchant_ref":   "MERCHANT_REF",
    "descriptor":     "DESCRIPTOR",
    "ip":             "IP",
    "ip_location":    "IP_LOCATION",
    "access_time":    "ACCESS_TIME",
    "device":         "DEVICE",
    "device_id":      "DEVICE_ID",
}

def load_config():
    # GitHub Actions: read from environment variables set by the workflow
    if os.environ.get("GITHUB_ACTIONS"):
        cfg = {}
        for key, env in ENV_MAP.items():
            val = os.environ.get(env, "").strip()
            cfg[key] = val if val and val != "N/A" else DEFAULTS[key]
        return cfg

    # Local: read config.json, fall back to defaults
    cfg_path = BASE / "config.json"
    if cfg_path.exists():
        with open(cfg_path, encoding="utf-8") as f:
            return {**DEFAULTS, **json.load(f)}
    print("[INFO] config.json not found — using defaults.")
    return DEFAULTS.copy()

# ── Page / margin helpers ──────────────────────────────────────────────────────
TEXT_WIDTH_CM = 15.92  # A4 minus 2 × 2.54 cm margins

def set_margins(doc):
    for s in doc.sections:
        s.page_width  = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = s.right_margin = Cm(2.54)
        s.top_margin  = s.bottom_margin = Cm(2.54)

# ── Paragraph helpers ──────────────────────────────────────────────────────────
def _para(doc, text, bold=False, size=11, center=False,
          indent_cm=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    r = p.add_run(text)
    r.bold       = bold
    r.font.size  = Pt(size)
    r.font.name  = "Calibri"
    return p

def title(doc, text):
    return _para(doc, text, bold=True, size=16, center=True,
                 space_before=12, space_after=12)

def section_head(doc, text):
    return _para(doc, text, bold=True, size=11,
                 space_before=8, space_after=4)

def body(doc, text, indent_cm=0, space_after=3):
    return _para(doc, text, size=11, indent_cm=indent_cm, space_after=space_after)

def bold_label(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(11); r2.font.name = "Calibri"
    return p

def spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r = OxmlElement("w:r")
    r.append(br)
    p._p.append(r)
    return p

# ── Image helpers ──────────────────────────────────────────────────────────────
def add_image(doc, img_bytes):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_picture(BytesIO(img_bytes), width=Cm(TEXT_WIDTH_CM))
    return p

def render_pdf_page(fitz_page, scale=3.0):
    pix = fitz_page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
    return pix.tobytes("png")

def load_file_as_image(path):
    path = Path(path)
    if not path.exists():
        return None
    if path.suffix.lower() == ".pdf":
        doc = fitz.open(str(path))
        return render_pdf_page(doc[0])
    return path.read_bytes()

def find(folder, names, exts=(".pdf", ".png", ".jpg", ".jpeg")):
    for name in names:
        for ext in exts:
            p = folder / (name + ext)
            if p.exists():
                return p
    return None

# ── PDF text extraction ────────────────────────────────────────────────────────
def extract_pages(path):
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(x_tolerance=2, y_tolerance=3)
            pages.append(txt.strip() if txt else "")
    return pages

def classify(line):
    if re.match(r"^(Checkout Flow|Refund & Cancellation Policy|Terms of Service)\s*$", line):
        return "title"
    if re.match(r"^Section\s+\d+[\s\u2013\-]", line):
        return "section"
    if re.match(r"^(Overview|Helpful Links:|IN ACCORDANCE WITH APPLICABLE LAWS)", line):
        return "section"
    if re.match(r"^\d{1,2}\.\d+(\.\d+)*\.?\s+\S", line):
        return "indent"
    if re.match(r"^(18|19)\.\d+\.?\s+\S", line):
        return "indent"
    if line.lstrip().startswith("\u2022"):
        return "indent"
    return "body"

def add_text_doc(doc, pages):
    for pi, page_text in enumerate(pages):
        if pi > 0:
            page_break(doc)
        for raw in page_text.split("\n"):
            raw = raw.strip()
            if not raw:
                continue
            kind = classify(raw)
            if kind == "title":
                title(doc, raw)
            elif kind == "section":
                section_head(doc, raw)
            elif kind == "indent":
                body(doc, raw, indent_cm=0.9)
            else:
                body(doc, raw)

# ── Main ──────────────────────────────────────────────────────────────────────
def generate():
    cfg = load_config()

    doc = Document()
    # Remove the default empty paragraph python-docx inserts
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    set_margins(doc)

    # ── 1. Checkout Flow ──────────────────────────────────────────────────────
    cf = find(PDFS, ["checkout_flow", "Checkout_Flow", "checkout"])
    if cf:
        print(f"[OK] Checkout Flow: {cf.name}")
        add_text_doc(doc, extract_pages(cf))
    else:
        print("[SKIP] checkout_flow.pdf not found")
        title(doc, "Checkout Flow")
        body(doc, "[Upload checkout_flow.pdf to the pdfs/ folder]")

    # ── 2. Refund & Cancellation Policy ───────────────────────────────────────
    page_break(doc)
    rp = find(PDFS, ["refund_cancellation", "refund", "Refund_Cancellation",
                      "Refund_&_Cancellation_Policy"])
    if rp:
        print(f"[OK] Refund Policy: {rp.name}")
        add_text_doc(doc, extract_pages(rp))
    else:
        print("[SKIP] refund_cancellation.pdf not found")
        title(doc, "Refund & Cancellation Policy")
        body(doc, "[Upload refund_cancellation.pdf to the pdfs/ folder]")

    # ── 3. Terms of Service ───────────────────────────────────────────────────
    page_break(doc)
    tos = find(PDFS, ["terms_of_service", "Terms_of_Service", "terms", "tos"])
    if tos:
        print(f"[OK] Terms of Service: {tos.name}")
        add_text_doc(doc, extract_pages(tos))
    else:
        print("[SKIP] terms_of_service.pdf not found")
        title(doc, "Terms of Service")
        body(doc, "[Upload terms_of_service.pdf to the pdfs/ folder]")

    # ── 4. Invoice ────────────────────────────────────────────────────────────
    page_break(doc)
    _para(doc, "Invoice:", bold=True, size=12, space_after=6)
    inv = find(PDFS, ["invoice", "Invoice"])
    if inv:
        print(f"[OK] Invoice: {inv.name}")
        add_image(doc, load_file_as_image(inv))
    else:
        print("[SKIP] invoice.pdf/png — upload to pdfs/")
        body(doc, "[Upload invoice.pdf or invoice.png to pdfs/]")

    # ── 5. Customer Communication ─────────────────────────────────────────────
    page_break(doc)
    _para(doc, "Customer Communication proof with our support channel:",
          bold=True, size=12, space_after=6)
    ic = find(PDFS, ["intercom", "customer_communication", "customer_email"])
    if ic:
        print(f"[OK] Intercom: {ic.name}")
        add_image(doc, load_file_as_image(ic))
    else:
        print("[SKIP] intercom.png/pdf — upload to pdfs/")
        body(doc, "[Upload intercom.png or intercom.pdf to pdfs/]")

    spacer(doc)
    body(doc,
         "We acknowledge the customer\u2019s dispute and would like to provide a comprehensive "
         "response to the chargeback filed. We maintain that the transaction in question was "
         "processed accurately and in full compliance with our terms and conditions.",
         space_after=6)
    body(doc,
         "Furthermore, the customer did not raise any concerns or discrepancies with us prior "
         "to initiating the dispute. We have established multiple support channels to address "
         "any issues or inquiries that our customers may have, yet no communication was received "
         "from the customer regarding this matter.",
         space_after=6)

    # ── 6. Proof of Service Delivery ──────────────────────────────────────────
    page_break(doc)
    _para(doc, "Proof of Service Delivery:", bold=True, size=12, space_after=6)
    gm = find(PDFS, ["gmail", "service_delivery", "confirmation_email"])
    if gm:
        print(f"[OK] Gmail: {gm.name}")
        add_image(doc, load_file_as_image(gm))
    else:
        print("[SKIP] gmail.png/pdf — upload to pdfs/")
        body(doc, "[Upload gmail.png or gmail.pdf to pdfs/]")

    # ── 7. Customer Purchase Details + Admin ──────────────────────────────────
    page_break(doc)
    body(doc,
         f"On {cfg['purchase_date']}, a customer named \u2013 {cfg['customer_name']}, "
         f"using email {cfg['customer_email']}, took our service for "
         f"${cfg['amount']} using their {cfg['card_type']} card. "
         f"The customer took the service without any influence.")
    spacer(doc)
    _para(doc, "The payment details of this customer are as follows:",
          bold=True, size=11, space_after=4)
    bold_label(doc, "Merchant Payment Reference", cfg["merchant_ref"])
    bold_label(doc, "Card", f"{cfg['card_type']} ending {cfg['card_number']}")
    bold_label(doc, "Amount", f"${cfg['amount']}")
    bold_label(doc, "Descriptor", cfg["descriptor"])
    spacer(doc)
    body(doc,
         "It\u2019s important to highlight that the client not only acquired and utilized the "
         "product but also contravened the Terms and Conditions (T&C) of the agreement "
         "post-purchase. Such a sequence of events undermines the claim that the client did "
         "not receive the product, rendering it implausible.")
    spacer(doc)
    _para(doc, "Customer Purchase & Breaching of (T&C) Proof:",
          bold=True, size=11, space_after=4)
    adm = find(PDFS, ["admin", "admin_panel", "tc_proof", "admin_screenshot"])
    if adm:
        print(f"[OK] Admin: {adm.name}")
        add_image(doc, load_file_as_image(adm))
    else:
        print("[SKIP] admin.png/pdf — upload to pdfs/")
        body(doc, "[Upload admin.png or admin.pdf to pdfs/]")

    # ── 8. Daily Drawdown ─────────────────────────────────────────────────────
    page_break(doc)
    _para(doc, "What is the Maximum Daily Drawdown Limit?",
          bold=True, size=12, space_after=8)
    _para(doc, "For Evaluation, Express, and Stellar 2-Step:",
          bold=True, size=11, space_after=4)
    body(doc,
         "Traders are allowed to lose 5% of their initial account balance on any given day "
         "as their daily drawdown. If your initial account balance is $100,000, then 5% of "
         "the amount is $5,000. At any point in a day if you are losing more than $5,000, "
         "your account will be automatically closed.")
    _para(doc, "For Stellar 1-Step:", bold=True, size=11,
          space_before=6, space_after=4)
    body(doc,
         "Traders are allowed to lose 3% of their initial account balance on any given day "
         "as their daily drawdown.")
    ddc = find(PDFS, ["drawdown_chart", "drawdown", "chart"])
    if ddc:
        print(f"[OK] Drawdown chart: {ddc.name}")
        add_image(doc, load_file_as_image(ddc))

    # ── 9. Customer IP and Access Log ─────────────────────────────────────────
    page_break(doc)
    _para(doc, "Customer IP and Access Log", bold=True, size=12, space_after=8)
    body(doc,
         "We acknowledge the customer\u2019s assertion that they have not received the product "
         "or service they purchased from us. However, we possess substantial evidence that "
         "contradicts this claim. Our records include access logs, IP tracking data, and device "
         "identification information that unequivocally demonstrate the customer\u2019s active "
         "engagement with our platform following the purchase.")
    spacer(doc)
    body(doc,
         "This data serves as irrefutable evidence that the customer not only received but also "
         "actively used our product, thereby invalidating their chargeback claim.")
    spacer(doc)
    _para(doc, "Customer Access Details:", bold=True, size=11, space_after=4)
    for label, key in [
        ("Name",               "customer_name"),
        ("Email",              "customer_email"),
        ("Account Number",     "account_login"),
        ("IP Address",         "ip"),
        ("IP Location",        "ip_location"),
        ("Access Date & Time", "access_time"),
        ("Customer Device",    "device"),
        ("Device ID",          "device_id"),
    ]:
        bold_label(doc, label, cfg.get(key, "N/A"))

    # ── 10. Journal ───────────────────────────────────────────────────────────
    jf = find(PDFS, ["journal", "mt5_journal", "trading_history", "manager_journal"])
    if jf and jf.suffix.lower() == ".pdf":
        print(f"[OK] Journal: {jf.name}")
        jdoc = fitz.open(str(jf))
        page_break(doc)
        _para(doc, "Trading History from Manager", bold=True, size=12, space_after=8)
        for pi in range(len(jdoc)):
            if pi > 0:
                page_break(doc)
            add_image(doc, render_pdf_page(jdoc[pi], scale=3.0))
        jdoc.close()
    else:
        print("[SKIP] journal.pdf — upload to pdfs/")
        page_break(doc)
        _para(doc, "Trading History from Manager", bold=True, size=12, space_after=8)
        body(doc, "[Upload journal.pdf to the pdfs/ folder]")

    # ── Save ──────────────────────────────────────────────────────────────────
    safe_arn  = re.sub(r"[^a-zA-Z0-9]", "_", cfg.get("arn", "dispute"))[:24]
    safe_date = cfg.get("purchase_date", "").replace("-", "")
    out = BASE / f"FN_Dispute_{safe_arn}_{safe_date}.docx"
    doc.save(str(out))
    print(f"\n\u2713 Done: {out.name}")


if __name__ == "__main__":
    generate()
